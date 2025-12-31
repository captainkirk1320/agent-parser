from io import BytesIO
from docx import Document
from fastapi.testclient import TestClient
from app.main import app

client = TestClient(app)

def test_parse_docx_extracts_email():
    doc = Document()
    doc.add_paragraph("Jane Doe")
    doc.add_paragraph("jane.doe@example.com")
    doc.add_paragraph("(555) 123-4567")
    doc.add_paragraph("Skills: Python, FastAPI")

    buf = BytesIO()
    doc.save(buf)

    files = {
        "file": ("resume.docx", buf.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    }
    r = client.post("/parse", files=files)
    assert r.status_code == 200
    data = r.json()

    assert data["candidate_profile"]["email"] == "jane.doe@example.com"
    assert len(data["evidence_map"]["email"]) >= 1
    assert data["evidence_map"]["email"][0]["locator"].startswith("docx:paragraph:")
